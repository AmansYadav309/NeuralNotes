import os
import hashlib
import uuid
import pdfplumber
from typing import List, Optional, Dict, Any



from fastapi import FastAPI, File, UploadFile, HTTPException, Body, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.concurrency import run_in_threadpool

from pydantic import BaseModel

import asyncio
from dotenv import load_dotenv

import httpx

load_dotenv()

# Environment variables
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
FRONTEND_ORIGIN = os.getenv("FRONTEND_ORIGIN", "http://localhost:3000")

app = FastAPI(title="NeuralNotes Backend", version="1.1.0")

def _ensure_dirs():
    os.makedirs("uploads", exist_ok=True)
    os.makedirs("media", exist_ok=True)

app.mount("/media", StaticFiles(directory="media"), name="media")


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

_ensure_dirs()


OPENROUTER_BASE = "https://openrouter.ai/api/v1"

class UploadResponse(BaseModel):
    num_pages: int
    pages: List[str]
    metadata: Dict[str, Any] = {}

class SummarizeRequest(BaseModel):
    page_number: int
    text: str
    format: Optional[str] = "paragraph"
    is_eli5: Optional[bool] = False

class SummarizeYouTubeRequest(BaseModel):
    url: str
    format: Optional[str] = "paragraph"
    is_eli5: Optional[bool] = False

class TTSRequest(BaseModel):
    summary: str

class DoubtRequest(BaseModel):
    question: str
    context: str

class DoubtGeneralRequest(BaseModel):
    question: str

class ConceptMapRequest(BaseModel):
    context: str

class QuizRequest(BaseModel):
    text: str

class FlashcardRequest(BaseModel):
    text: str

class ExportRequest(BaseModel):
    format: str
    content_type: str
    content: Any
    title: str

class ExportGDocRequest(BaseModel):
    auth_code: str
    content_type: str
    content: Any
    title: str

class DiagramRequest(BaseModel):
    content: str
    diagram_type: str
    topic: str

# In-memory caches
SUMMARY_CACHE: Dict[str, Any] = {}

def _hash_key(prefix: str, *parts: str) -> str:
    h = hashlib.sha256()
    for p in parts:
        h.update(p.encode("utf-8"))
    return f"{prefix}:{h.hexdigest()}"

@app.post("/upload_pdf", response_model=UploadResponse)
async def upload_pdf(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported.")

    path = os.path.join("uploads", f"{uuid.uuid4()}_{file.filename}")

    try:
        with open(path, "wb") as f:
            f.write(await file.read())
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save upload: {e}")

    pages_text: List[str] = []
    meta: Dict[str, Any] = {}
    try:
        with pdfplumber.open(path) as pdf:
            meta = pdf.metadata or {}
            for page in pdf.pages:
                txt = page.extract_text() or ""
                pages_text.append(txt)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read PDF: {e}")

    return UploadResponse(num_pages=len(pages_text), pages=pages_text, metadata=meta)

async def _openrouter_chat(system_prompt: str, user_prompt: str, model: str = "google/gemini-2.0-flash-001") -> str:
    if not OPENROUTER_API_KEY:
        raise HTTPException(status_code=500, detail="OPENROUTER_API_KEY not configured")
    
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "HTTP-Referer": FRONTEND_ORIGIN,
        "X-Title": "NeuralNotes"
    }
    
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": 0.3,
        "max_tokens": 8192
    }
    
    try:
        async with httpx.AsyncClient(timeout=120.0) as ac:
            res = await ac.post(f"{OPENROUTER_BASE}/chat/completions", headers=headers, json=payload)
            
            if res.status_code == 429:
                raise HTTPException(status_code=429, detail="AI Rate limit reached via OpenRouter. Please wait a moment.")
            
            res.raise_for_status()
            data = res.json()
            return data["choices"][0]["message"]["content"]
    except httpx.HTTPStatusError as e:
        error_info = e.response.text
        raise HTTPException(status_code=e.response.status_code, detail=f"OpenRouter API error: {error_info}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to communicate with OpenRouter: {e}")

@app.post("/summarize")
async def summarize(req: SummarizeRequest):
    key = _hash_key("sum", str(req.page_number), req.format, str(req.is_eli5), req.text)
    if key in SUMMARY_CACHE:
        return {"summary": SUMMARY_CACHE[key], "format": req.format, "cached": True}

    if req.format == "bullets":
        system = "You are an expert academic assistant."
        user = (
            "Summarize the following text as concise bullet points, one idea per bullet.\n\n"
            + req.text
        )
    elif req.format == "mindmap":
        system = "You are an expert academic assistant."
        user = (
            "Organize the following text into a hierarchical mind map outline with main topics and subtopics.\n\n"
            + req.text
        )
    elif req.format == "flashcards":
        system = "You are an expert academic assistant."
        user = (
            "Extract key terms and definitions from the following text as a list of flashcard pairs. "
            "You MUST return ONLY valid JSON formatted exactly as: `{\"flashcards\": [{\"term\": \"...\", \"definition\": \"...\"}]}`\n\n"
            + req.text
        )
    elif req.format == "mixed":
        system = "You are an expert academic assistant."
        user = (
            "First, provide a clear paragraph summary of the following text. "
            "Then, provide a few concise bullet points highlighting the key ideas, formulas, and definitions.\n\n"
            + req.text
        )
    else:
        system = "You are an expert academic assistant."
        user = (
            "Write a clear paragraph summary of the following text, focusing on key ideas, formulas, and definitions suitable for revision.\n\n"
            + req.text
        )
        
    if req.is_eli5:
        user += "\n\nCRITICAL INSTRUCTION: Explain this as if I am 5 years old. Use extremely simple analogies, basic vocabulary, and short sentences. Avoid all complex jargon."
    
    summary = await _openrouter_chat(system, user)
    
    if req.format == "flashcards":
        import json
        cleaned = summary.strip()
        if cleaned.startswith("```json"):
            cleaned = cleaned[7:]
        elif cleaned.startswith("```"):
            cleaned = cleaned[3:]
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]
        cleaned = cleaned.strip()
        try:
            parsed = json.loads(cleaned)
            summary_out = parsed.get("flashcards", [])
        except json.JSONDecodeError:
            summary_out = []
    else:
        summary_out = summary

    SUMMARY_CACHE[key] = summary_out
    return {"summary": summary_out, "format": req.format, "cached": False}

@app.post("/summarize-youtube")
async def summarize_youtube(req: SummarizeYouTubeRequest):
    try:
        from transcriber import download_youtube_audio, transcribe_audio, get_youtube_transcript
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Transcriber module error: {e}")
        
    try:
        # 1. Attempt instant transcription via YouTube API
        try:
            transcript = await run_in_threadpool(get_youtube_transcript, req.url)
            print("Successfully extracted instant YouTube transcript!")
        except Exception as transcript_err:
            print(f"Transcript API failed: {transcript_err}. Falling back to downloading and transcribing audio...")
            # 2. Download audio on a background thread
            audio_path = await run_in_threadpool(download_youtube_audio, req.url)
            # 3. Transcribe on a background thread
            transcript = await run_in_threadpool(transcribe_audio, audio_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to process YouTube video: {e}")
        
    # 3. Summarize using existing OpenRouter logic
    system = "You are an expert academic assistant."
    
    if req.format == "bullets":
        user_prompt_base = "Summarize this lecture transcript as concise bullet points, one idea per bullet.\n\n"
    elif req.format == "mindmap":
        user_prompt_base = "Organize this lecture transcript into a hierarchical mind map outline with main topics and subtopics.\n\n"
    elif req.format == "flashcards":
        user_prompt_base = (
            "Extract key terms and definitions from this lecture transcript as a list of flashcard pairs. "
            "You MUST return ONLY valid JSON formatted exactly as: `{\"flashcards\": [{\"term\": \"...\", \"definition\": \"...\"}]}`\n\n"
        )
    elif req.format == "mixed":
        user_prompt_base = "First, provide a clear paragraph summary of this lecture transcript. Then, provide a few concise bullet points highlighting the key ideas, concepts, and definitions.\n\n"
    else:
        user_prompt_base = "Summarize this lecture transcript clearly and concisely in a paragraph format, focusing on key ideas, concepts, and definitions suitable for revision.\n\n"
        
    user = user_prompt_base + transcript
    
    if req.is_eli5:
        user += "\n\nCRITICAL INSTRUCTION: Explain this as if I am 5 years old. Use extremely simple analogies, basic vocabulary, and short sentences. Avoid all complex jargon."
    
    summary = await _openrouter_chat(system, user)
    return {"summary": summary, "transcript": transcript}

@app.post("/summarize-video-file")
async def summarize_video_file(
    file: UploadFile = File(...),
    is_eli5: bool = Form(False),
    format: str = Form("paragraph")
):
    if not file.filename.lower().endswith(('.mp4', '.mp3', '.wav', '.m4a', '.weba')):
        raise HTTPException(status_code=400, detail="Unsupported audio/video format.")
        
    try:
        from transcriber import transcribe_audio, TEMP_DIR
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Transcriber module error: {e}")
        
    path = os.path.join(TEMP_DIR, f"{uuid.uuid4()}_{file.filename}")
    try:
        with open(path, "wb") as f:
            f.write(await file.read())
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save upload: {e}")
        
    try:
        # Transcribe directly on a background thread
        transcript = await run_in_threadpool(transcribe_audio, path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to transcribe media file: {e}")
        
    # Summarize transcript
    system = "You are an expert academic assistant."
    
    if format == "bullets":
        user_prompt_base = "Summarize this lecture transcript as concise bullet points, one idea per bullet.\n\n"
    elif format == "mindmap":
        user_prompt_base = "Organize this lecture transcript into a hierarchical mind map outline with main topics and subtopics.\n\n"
    elif format == "flashcards":
        user_prompt_base = (
            "Extract key terms and definitions from this lecture transcript as a list of flashcard pairs. "
            "You MUST return ONLY valid JSON formatted exactly as: `{\"flashcards\": [{\"term\": \"...\", \"definition\": \"...\"}]}`\n\n"
        )
    elif format == "mixed":
        user_prompt_base = "First, provide a clear paragraph summary of this lecture transcript. Then, provide a few concise bullet points highlighting the key ideas, concepts, and definitions.\n\n"
    else:
        user_prompt_base = "Summarize this lecture transcript clearly and concisely in a paragraph format, focusing on key ideas, concepts, and definitions suitable for revision.\n\n"
        
    user = user_prompt_base + transcript
    
    if is_eli5:
        user += "\n\nCRITICAL INSTRUCTION: Explain this as if I am 5 years old. Use extremely simple analogies, basic vocabulary, and short sentences. Avoid all complex jargon."
    
    summary = await _openrouter_chat(system, user)
    return {"summary": summary, "transcript": transcript}

@app.post("/tts")
async def tts(req: TTSRequest):
    try:
        from gtts import gTTS
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"gTTS not installed: {e}")

    if not req.summary.strip():
        raise HTTPException(status_code=400, detail="Summary is empty")
    filename = f"tts_{uuid.uuid4()}.mp3"
    filepath = os.path.join("media", filename)
    try:
        tts = gTTS(text=req.summary)
        tts.save(filepath)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"TTS failed: {e}")

    return {"url": f"/media/{filename}"}

@app.post("/doubt")
async def doubt(req: DoubtRequest):
    system = (
        "You are an expert academic tutor helping a student understand their study material.\n"
        "First, check if the provided STUDY MATERIAL contains information relevant to the STUDENT QUESTION.\n"
        "If yes:\n"
        "1. Set `found_in_doc` to true.\n"
        "2. Formulate your answer based ONLY on the provided context.\n"
        "3. Identify the page where you found the answer. The text is separated by `[PAGE N]` markers.\n"
        "4. Include a short excerpt (first ~10 words) of the specific paragraph.\n"
        "If no:\n"
        "1. Set `found_in_doc` to false.\n"
        "2. Leave `answer` blank.\n"
        "\n"
        "You MUST return ONLY valid JSON formatted exactly like this:\n"
        "{\n"
        '  "found_in_doc": true/false,\n'
        '  "answer": "Your detailed explanation here (if found)",\n'
        '  "source": {\n'
        '    "page": N,\n'
        '    "paragraph": 1,\n'
        '    "excerpt": "..."\n'
        '  }\n'
        "}"
    )
    user = (
        f"STUDY MATERIAL (FULL CONTEXT):\n{req.context}\n\n"
        f"STUDENT QUESTION:\n{req.question}\n\n"
    )
    
    response_text = await _openrouter_chat(system, user)
    
    # Strip markdown formatting if any
    if response_text.startswith("```json"):
        response_text = response_text[7:]
    elif response_text.startswith("```"):
        response_text = response_text[3:]
    if response_text.endswith("```"):
        response_text = response_text[:-3]
        
    import json
    try:
        data = json.loads(response_text.strip())
        return data
    except json.JSONDecodeError:
        print(f"Failed to parse JSON: {response_text}")
        raise HTTPException(status_code=500, detail="Failed to parse structured JSON from AI model.")

@app.post("/doubt-general")
async def doubt_general(req: DoubtGeneralRequest):
    system = "You are a helpful, expert AI tutor. Please explain the requested topic thoroughly."
    user = f"STUDENT QUESTION:\n{req.question}\n\nPlease provide a detailed, educational response."
    
    answer = await _openrouter_chat(system, user)
    return {
        "found_in_doc": False,
        "answer": answer,
        "source": "general_knowledge"
    }

@app.post("/concept-map")
async def concept_map(req: ConceptMapRequest):
    system = (
        "You are an expert curriculum designer. Extract a structured concept map out of the provided document. "
        "The concept map should go 2-3 levels deep and be organized logically. "
        "Each node should define a single topic, and include a 1-2 sentence summary.\n\n"
        "You MUST return ONLY valid JSON formatted exactly as this schema:\n"
        "{\n"
        '  "title": "Main Topic",\n'
        '  "children": [\n'
        "    {\n"
        '      "label": "Subtopic A",\n'
        '      "summary": "Brief explanation of Subtopic A",\n'
        '      "children": []\n'
        "    }\n"
        "  ]\n"
        "}"
    )
    user = f"DOCUMENT CONTENT:\n\n{req.context}"
    
    response_text = await _openrouter_chat(system, user)
    
    if response_text.startswith("```json"):
        response_text = response_text[7:]
    elif response_text.startswith("```"):
        response_text = response_text[3:]
    if response_text.endswith("```"):
        response_text = response_text[:-3]

    import json
    try:
        data = json.loads(response_text.strip())
        return data
    except json.JSONDecodeError:
        print(f"JSON Parse Error. LLM returned: {response_text}")
        raise HTTPException(status_code=500, detail="Failed to parse JSON tree from AI model.")

@app.post("/generate-quiz")
async def generate_quiz(req: QuizRequest):
    system = (
        "You are an expert educational AI. Your task is to generate a multiple-choice quiz based on the provided text. "
        "You MUST return ONLY valid JSON. Do not include any markdown formatting, headers, or conversational text. "
        "The JSON structure must match this EXACT format:\n"
        '{"questions": [{"question": "...", "options": ["Descriptive Option 1", "Descriptive Option 2", "Descriptive Option 3", "Descriptive Option 4"], "answer": "Exact text of correct option here", "explanation": "Why this is the correct answer"}]}'
    )
    user = f"Generate a 5-question quiz for the following text:\n\n{req.text}"
    
    response_text = await _openrouter_chat(system, user)
    
    # Strip markdown code fences if Gemini added them
    if response_text.startswith("```json"):
        response_text = response_text[7:]
    if response_text.startswith("```"):
        response_text = response_text[3:]
    if response_text.endswith("```"):
        response_text = response_text[:-3]
        
    import json
    try:
        data = json.loads(response_text.strip())
        return data
    except json.JSONDecodeError:
        print(f"Failed to parse JSON: {response_text}")
        raise HTTPException(status_code=500, detail="Failed to parse quiz JSON from AI model.")

@app.post("/generate-flashcards")
async def generate_flashcards(req: FlashcardRequest):
    system = (
        "You are an expert educational AI. Your task is to extract key terms and concepts from the provided text and format them as flashcards. "
        "You MUST return ONLY valid JSON. Do not include any markdown formatting, headers, or conversational text. "
        "The JSON structure must match this EXACT format:\n"
        '{"flashcards": [{"term": "...", "definition": "..."}]}'
    )
    user = f"Extract 5-10 key terms and definitions as flashcards from the following text:\n\n{req.text}"
    
    response_text = await _openrouter_chat(system, user)
    
    # Strip markdown code fences if Gemini added them
    if response_text.startswith("```json"):
        response_text = response_text[7:]
    if response_text.startswith("```"):
        response_text = response_text[3:]
    if response_text.endswith("```"):
        response_text = response_text[:-3]
        
    import json
    try:
        data = json.loads(response_text.strip())
        return data
    except json.JSONDecodeError:
        print(f"Failed to parse JSON: {response_text}")
        raise HTTPException(status_code=500, detail="Failed to parse flashcard JSON from AI model.")

@app.post("/export")
async def export_content(req: ExportRequest):
    format = req.format.lower()
    title = req.title
    content_type = req.content_type
    content = req.content
    
    filename = f"neuralnotes-export.{format}"
    docs_text = ""
    
    try:
        if content_type == "summary":
            if isinstance(content, str):
                docs_text = content
            elif isinstance(content, list):
                docs_text = "\n\n".join([str(c) for c in content if c])
            else:
                docs_text = str(content)
        elif content_type == "qa":
            q = content.get('question', '') if isinstance(content, dict) else 'Question'
            a = content.get('answer', '') if isinstance(content, dict) else str(content)
            docs_text = f"Q: {q}\n\nA: {a}"
        elif content_type == "quiz":
            docs_text = "Quiz\n"
            if isinstance(content, list):
                for i, q in enumerate(content):
                    docs_text += f"\nQ{i+1}: {q.get('question', '')}\n"
                    # Add options
                    opts = q.get('options', [])
                    for opt in opts:
                        docs_text += f" - {opt}\n"
                    docs_text += f"Answer: {q.get('answer', '')}\n"
                    docs_text += f"Explanation: {q.get('explanation', '')}\n"
        elif content_type == "flashcards":
            docs_text = "Flashcards\n"
            if isinstance(content, list):
                for i, f in enumerate(content):
                    docs_text += f"\nCard {i+1}:\nFront: {f.get('term', '')}\nBack: {f.get('definition', '')}\n"
    except Exception as e:
        docs_text = str(content)

    temp_path = os.path.join("media", f"{uuid.uuid4()}-{filename}")
    
    if format == "pdf":
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        # FPDF default Helvetica might not handle all Unicode
        safe_title = title.encode('latin-1', 'replace').decode('latin-1')
        safe_docs_text = docs_text.encode('latin-1', 'replace').decode('latin-1')
        
        pdf.cell(0, 10, txt=safe_title, ln=1, align='C')
        pdf.ln(10)
        pdf.multi_cell(0, 10, txt=safe_docs_text)
        pdf.output(temp_path)
        return FileResponse(path=temp_path, filename=filename, media_type="application/pdf")

    elif format == "docx":
        from docx import Document
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(docs_text)
        doc.save(temp_path)
        return FileResponse(path=temp_path, filename=filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
    elif format == "txt":
        out = f"{title}\n\n{docs_text}"
        return Response(content=out, media_type="text/plain", headers={"Content-Disposition": f'attachment; filename="{filename}"'})

@app.post("/export-gdoc")
async def export_gdoc(req: ExportGDocRequest):
    try:
        from google_auth_oauthlib.flow import Flow
        from googleapiclient.discovery import build
    except ImportError:
        raise HTTPException(status_code=500, detail="Google API client libraries not installed.")

    creds_path = os.getenv("GOOGLE_CREDENTIALS_PATH", "./google-credentials.json")
    if not os.path.exists(creds_path):
        raise HTTPException(status_code=500, detail="Google credentials file not found.")

    try:
        flow = Flow.from_client_secrets_file(
            creds_path,
            scopes=["https://www.googleapis.com/auth/documents", "https://www.googleapis.com/auth/drive.file"],
            redirect_uri="postmessage"
        )
        flow.fetch_token(code=req.auth_code)
        creds = flow.credentials
    except Exception as e:
        print(f"OAuth Error: {e}")
        raise HTTPException(status_code=401, detail=f"Authentication failed: {str(e)}")

    docs_text = ""
    # Process content identically to /export
    try:
        content_type = req.content_type
        content = req.content
        if content_type == "summary":
            if isinstance(content, str):
                docs_text = content
            elif isinstance(content, list):
                docs_text = "\n\n".join([str(c) for c in content if c])
            else:
                docs_text = str(content)
        elif content_type == "qa":
            q = content.get('question', '') if isinstance(content, dict) else 'Question'
            a = content.get('answer', '') if isinstance(content, dict) else str(content)
            docs_text = f"Q: {q}\n\nA: {a}"
        elif content_type == "quiz":
            docs_text = "Quiz\n"
            if isinstance(content, list):
                for i, q in enumerate(content):
                    docs_text += f"\nQ{i+1}: {q.get('question', '')}\n"
                    opts = q.get('options', [])
                    for opt in opts:
                        docs_text += f" - {opt}\n"
                    docs_text += f"Answer: {q.get('answer', '')}\n"
                    docs_text += f"Explanation: {q.get('explanation', '')}\n"
        elif content_type == "flashcards":
            docs_text = "Flashcards\n"
            if isinstance(content, list):
                for i, f in enumerate(content):
                    docs_text += f"\nCard {i+1}:\nFront: {f.get('term', '')}\nBack: {f.get('definition', '')}\n"
    except Exception as e:
        docs_text = str(req.content)

    try:
        service = build('docs', 'v1', credentials=creds)
        
        # Create a blank document
        document = service.documents().create(body={'title': req.title}).execute()
        document_id = document.get('documentId')
        
        # Insert text
        full_text = f"{req.title}\n\n{docs_text}"
        requests = [
            {
                'insertText': {
                    'location': {
                        'index': 1,
                    },
                    'text': full_text
                }
            }
        ]
        
        service.documents().batchUpdate(documentId=document_id, body={'requests': requests}).execute()
        
        url = f"https://docs.google.com/document/d/{document_id}/edit"
        return {"url": url}
    except Exception as e:
        print(f"Google Docs API Error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to create Google Doc: {str(e)}")

@app.post("/generate-diagram")
async def generate_diagram(req: DiagramRequest):
    if req.diagram_type == "flowchart":
        base_prompt = "You are a diagram generator for students. Based on the following study content, generate a valid Mermaid.js flowchart TD diagram that visually explains the key concepts, processes, and relationships. Output ONLY the raw Mermaid DSL code — no explanation, no markdown code fences, no backticks. Start directly with flowchart TD. Keep node labels under 5 words each. Maximum 12 nodes."
    elif req.diagram_type == "sequence":
        base_prompt = "You are a diagram generator for students. Based on the following study content, generate a valid Mermaid.js sequenceDiagram that visually explains the cause-and-effect chains or interactions between components. Output ONLY the raw Mermaid DSL code — no explanation, no markdown code fences, no backticks. Start directly with sequenceDiagram. Keep node labels under 5 words each. Maximum 12 interactions."
    elif req.diagram_type == "mindmap":
        base_prompt = "You are a diagram generator for students. Based on the following study content, generate a valid Mermaid.js mindmap diagram that visually explains the topic overview and hierarchy. Output ONLY the raw Mermaid DSL code — no explanation, no markdown code fences, no backticks. Start directly with mindmap. Keep node labels under 5 words each. Maximum 12 nodes."
    else:
        raise HTTPException(status_code=400, detail="Invalid diagram type")
        
    user_prompt = f"Topic: {req.topic}\n\nContent:\n{req.content}"
    
    response_text = await _openrouter_chat(base_prompt, user_prompt)
    
    if response_text.startswith("```mermaid"):
        response_text = response_text[10:]
    elif response_text.startswith("```"):
        response_text = response_text[3:]
    if response_text.endswith("```"):
        response_text = response_text[:-3]

    return {"mermaid_code": response_text.strip()}

async def extract_pdf_text_from_bytes(contents: bytes) -> str:
    import io
    import pdfplumber
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(contents)) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
    except Exception as e:
        print(f"Error extracting PDF: {e}")
    return text

@app.post("/analyze-papers")
async def analyze_papers(files: List[UploadFile] = File(...)):
    all_papers_text = []
    for i, file in enumerate(files):
        contents = await file.read()
        text = await extract_pdf_text_from_bytes(contents)
        all_papers_text.append(f"=== PAPER {i+1}: {file.filename} ===\n{text}")

    combined_text = "\n\n".join(all_papers_text)

    prompt = f"""
You are an expert exam paper analyzer for university students in India.

Analyze the following {len(files)} question papers carefully.

Your tasks:
1. Extract every question from every paper
2. Detect which UNIT each question belongs to (Unit 1 to Unit 5, or as labeled). Group everything without a unit into 'Unit 1' if you are unsure.
3. Find questions that are repeated or very similar across papers
4. Count frequency of each question/topic
5. Classify priority: HIGH (asked 3+ times or 2+ times in different papers), MEDIUM (asked 2 times), LOW (asked once)

Return ONLY valid JSON in exactly this structure:
{{
  "subject_name": "detected subject name",
  "total_papers_analyzed": {len(files)},
  "units": [
    {{
      "unit_number": 1,
      "unit_name": "detected unit name if visible",
      "questions": [
        {{
          "topic": "short topic name",
          "core_idea": "one sentence description",
          "frequency": 2,
          "priority": "HIGH",
          "question_variants": [
            "Exact question text from paper 1",
            "Similar question text from paper 2"
          ],
          "appeared_in_papers": [1, 2]
        }}
      ]
    }}
  ]
}}

Here are the question papers:

{combined_text}
"""
    response = await _openrouter_chat("You are an expert exam analyzer. You only output valid JSON.", prompt)

    import json, re
    # Strip markdown if needed
    cleaned = response.strip()
    if cleaned.startswith("```json"):
        cleaned = cleaned[7:]
    elif cleaned.startswith("```"):
        cleaned = cleaned[3:]
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3]
    cleaned = cleaned.strip()

    try:
        analysis = json.loads(cleaned)
        # Ensure units array exists
        if "units" not in analysis:
            analysis["units"] = []
        return analysis
    except json.JSONDecodeError as e:
        print(f"Failed to parse analysis JSON. Error: {e}\nResponse: {response}")
        # Try regex extract
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            try:
                return json.loads(json_match.group())
            except:
                pass
        raise HTTPException(status_code=500, detail="Could not parse analysis from AI. Please try again.")

@app.post("/generate-exam-answer")
async def generate_exam_answer(
    question: str = Form(...),
    topic: str = Form(...),
    marks: int = Form(default=8),
    notes_file: Optional[UploadFile] = File(default=None)
):
    notes_context = ""
    if notes_file and notes_file.filename:
        try:
            contents = await notes_file.read()
            notes_text = await extract_pdf_text_from_bytes(contents)
            notes_context = f"""
Use the following teacher notes as your primary source. Align your answer with this material:

{notes_text[:12000]}  # limit to avoid token overflow
"""
        except Exception as e:
            print(f"Could not read notes file: {e}")

    prompt = f"""
You are helping an Indian university student prepare for exams.

Generate a {marks}-mark exam answer for the following question.

Question: {question}
Topic: {topic}
{notes_context}

Answer format rules:
- Write in simple, clear Indian English
- Start with a 1-2 line definition
- Use Markdown formatting: bold text for headings like **Definition**, **Types**, **Importance**, **Example**
- Use numbered points or bullet points inside sections for readability
- End with a real-world example
- Total length: suitable for {marks} marks (roughly {marks * 30} words)
- Do NOT use complex jargon without explanation
- Use proper line breaks so the text is highly readable in a Markdown viewer.

Generate the detailed answer now:
"""

    answer = await _openrouter_chat("You are an expert exam tutor.", prompt)
    return {
        "question": question,
        "topic": topic,
        "marks": marks,
        "answer": answer
    }

@app.get("/")
async def root():
    return {"status": "ok", "message": "NeuralNotes API with OpenRouter is running"}

@app.get("/health")
async def health():
    return {"status": "ok"}
